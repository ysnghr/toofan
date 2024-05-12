from analyzers.file_analyzer import FileAnalyzer
import docx
import msoffcrypto
from oletools.olevba import VBA_Parser
from langdetect import detect


class OfficeAnalyzer(FileAnalyzer):
    """
    Class for analyzing Office files.
    """
    def analyze(self, file):
        info = {
            "language_code": self._get_language_code(file),
            "page_count": self._get_page_count(file),
            "encrypted": self._is_encrypted(file),
            "macros": self._has_macros(file),
        }
        return info

    def _get_language_code(self, file):
        try:
            doc = docx.Document(file)
            text = " ".join([para.text for para in doc.paragraphs if para.text.strip() != ''])
            return detect(text) if text else 'Unknown'
        except Exception as e:
            return "Unknown"

    def _is_encrypted(self, file):
        try:
            office_file = msoffcrypto.OfficeFile(open(file, "rb"))
            return office_file.is_encrypted()
        except Exception as e:
            return f"Error: {str(e)}"

    def _has_macros(self, file):
        vbaparser = VBA_Parser(file)
        if vbaparser.detect_vba_macros():
            return True
        return False

    def _get_page_count(self, file):
        try:
            doc = docx.Document(file)
            return len(doc.element.xpath('.//w:sectPr'))
        except Exception as e:
            return None